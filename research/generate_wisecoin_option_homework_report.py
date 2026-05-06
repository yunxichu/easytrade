"""
generate_wisecoin_option_homework_report.py
==========================================

Generate the Word/PDF/ZIP submission for the wisecoin_option homework.
This script intentionally depends only on document-generation libraries and
previously generated CSV/JSON outputs, so it can run with the bundled document
runtime.
"""

from __future__ import annotations

import json
import shutil
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_NAME = "23300180062_任昊来_第九周作业"
DATA_DIR = ROOT / "data/wisecoin_option"
OUTPUT_DIR = ROOT / "research/outputs/wisecoin_option"
PACKAGE_PARENT = ROOT / "作业提交包" / "_staging_wisecoin"
PACKAGE_DIR = PACKAGE_PARENT / SUBMISSION_NAME
REPORT_DOCX = PACKAGE_DIR / f"{SUBMISSION_NAME}.docx"
REPORT_PDF = PACKAGE_DIR / f"{SUBMISSION_NAME}.pdf"


def pct(x: float) -> str:
    if pd.isna(x):
        return ""
    return f"{float(x):.2%}"


def num(x: float, digits: int = 2) -> str:
    if pd.isna(x):
        return ""
    return f"{float(x):.{digits}f}"


def money(x: float) -> str:
    if pd.isna(x):
        return ""
    return f"{float(x):,.2f}"


def ensure_package_dir() -> None:
    if PACKAGE_DIR.exists():
        shutil.rmtree(PACKAGE_DIR)
    for sub in ["code", "data", "outputs"]:
        (PACKAGE_DIR / sub).mkdir(parents=True, exist_ok=True)


def copy_artifacts() -> None:
    files = [
        (ROOT / "research/wisecoin_option_analysis.py", PACKAGE_DIR / "code/wisecoin_option_analysis.py"),
        (
            ROOT / "research/generate_wisecoin_option_homework_report.py",
            PACKAGE_DIR / "code/generate_wisecoin_option_homework_report.py",
        ),
    ]
    files += [(path, PACKAGE_DIR / "data" / path.name) for path in DATA_DIR.glob("*.csv")]
    files += [(path, PACKAGE_DIR / "outputs" / path.name) for path in OUTPUT_DIR.glob("*")]
    for src, dst in files:
        if src.exists() and src.is_file():
            shutil.copy2(src, dst)


def write_readme() -> None:
    (PACKAGE_DIR / "README.txt").write_text(
        "\n".join(
            [
                "wisecoin_option程序分析与组合策略扩展作业包",
                "",
                f"提交名称：{SUBMISSION_NAME}",
                "",
                "主要内容：",
                "1. 分析wisecoin_option类程序的数据下载流程和观点形成逻辑。",
                "2. 新增模块A：组合策略希腊字母计算。",
                "3. 新增模块B：策略盈亏分析，包括",
                "   - 到期盈亏曲线（161点）",
                "   - 7天后价格×IV情景重估热力图",
                "   - 1/7/14/21天多时间断面情景",
                "   - 风险中性 + 真实世界双假设的Monte Carlo模拟（20000路径）",
                "     给出POP、期望盈亏、VaR95、CVaR95",
                "",
                "主要文件：",
                f"- {SUBMISSION_NAME}.docx：Word报告。",
                f"- {SUBMISSION_NAME}.pdf：PDF报告。",
                "- code/wisecoin_option_analysis.py：数据下载、观点形成、组合Greeks、所有盈亏分析。",
                "- code/generate_wisecoin_option_homework_report.py：生成报告和提交包。",
                "- data/：Yahoo/yfinance下载的SPY历史数据和期权链。",
                "- outputs/：策略腿、组合Greeks、盈亏表、观点摘要和图形（共6张）。",
                "",
                "复现命令：",
                "python research/wisecoin_option_analysis.py",
                "python research/generate_wisecoin_option_homework_report.py",
            ]
        ),
        encoding="utf-8",
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.5)
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, "1F4E79")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_para(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(18)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.3)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.8)


def load_outputs() -> dict[str, object]:
    view = json.loads((OUTPUT_DIR / "view_summary.json").read_text(encoding="utf-8"))
    legs = pd.read_csv(OUTPUT_DIR / "strategy_legs.csv")
    greeks = pd.read_csv(OUTPUT_DIR / "portfolio_greeks.csv")
    payoff = pd.read_csv(OUTPUT_DIR / "strategy_payoff.csv")
    scenario = pd.read_csv(OUTPUT_DIR / "scenario_pnl.csv")
    mc_stats = pd.read_csv(OUTPUT_DIR / "monte_carlo_stats.csv") if (OUTPUT_DIR / "monte_carlo_stats.csv").exists() else None
    multi_horizon = pd.read_csv(OUTPUT_DIR / "multi_horizon_pnl.csv") if (OUTPUT_DIR / "multi_horizon_pnl.csv").exists() else None
    return {
        "view": view,
        "legs": legs,
        "greeks": greeks,
        "payoff": payoff,
        "scenario": scenario,
        "mc_stats": mc_stats,
        "multi_horizon": multi_horizon,
    }


def build_docx() -> None:
    data = load_outputs()
    view = data["view"]
    legs = data["legs"]
    greeks = data["greeks"]
    scenario = data["scenario"]
    mc_stats = data["mc_stats"]
    multi_horizon = data["multi_horizon"]
    total = greeks[greeks["contract_symbol"] == "TOTAL"].iloc[0]
    pnl = view["pnl_summary"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("wisecoin_option程序分析、组合希腊字母与策略盈亏")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    meta = doc.add_paragraph(f"提交文件名：{SUBMISSION_NAME}    数据样本：{view['ticker']} Yahoo期权链")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)

    add_heading(doc, "一、程序审查：数据下载与观点形成")
    add_para(
        doc,
        "本地仓库未发现老师给出的原始 wisecoin_option 源码，因此本作业按该类程序的典型逻辑进行复现与扩展：先从 Yahoo Finance 下载标的历史价格与期权链，再用趋势、波动率和技术指标形成方向观点，最后把观点映射为期权组合策略。",
    )
    add_bullet(doc, f"数据下载：使用 yfinance 获取 {view['ticker']} 近一年日线价格，并选取 DTE 接近45天的到期月。本次实际数据源为 {view['data_source']}，样本日 {view['as_of']}，到期日 {view['expiry']}，DTE={view['dte']} 天。")
    add_bullet(doc, "清洗规则：期权价格优先使用 bid/ask 中间价；若买卖价不可用则使用 lastPrice；过滤价格或IV缺失的合约。")
    add_bullet(doc, f"观点形成：MA20={num(view['ma20'])}，MA60={num(view['ma60'])}，20日收益={pct(view['ret20'])}，RSI14={num(view['rsi14'])}，形成方向观点：{view['directional_view_cn']}。")
    add_bullet(doc, f"波动率判断：ATM IV={pct(view['atm_iv'])}，RV20={pct(view['rv20'])}，IV-RV20={pct(view['iv_minus_rv20'])}，结论为：{view['volatility_view']}。")

    if (OUTPUT_DIR / "underlying_view.png").exists():
        doc.add_picture(str(OUTPUT_DIR / "underlying_view.png"), width=Inches(5.7))
        cap = doc.add_paragraph("图1  标的价格与趋势指标")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "二、策略选择")
    add_para(
        doc,
        f"程序观点为{view['directional_view_cn']}，同时IV高于近期RV，单腿买权会承受较高权利金与Theta消耗，因此选择有限风险的{view['strategy_name']}。该组合用买入较低行权价看涨期权表达上行，同时卖出较高行权价看涨期权降低成本。",
    )
    table = doc.add_table(rows=1, cols=6)
    headers = ["合约", "类型", "行权价", "方向", "中间价", "IV"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for _, row in legs.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["contract_symbol"])
        cells[1].text = str(row["option_type"])
        cells[2].text = num(row["strike"], 0)
        cells[3].text = "买入" if int(row["quantity"]) > 0 else "卖出"
        cells[4].text = money(row["mid"])
        cells[5].text = pct(row["implied_volatility"])
    style_table(table)

    add_heading(doc, "三、模块A：组合策略希腊字母计算")
    add_para(
        doc,
        "新增模块A的核心是先用Black-Scholes模型计算每条期权腿的Delta、Gamma、Vega、Theta、Rho，再按持仓方向、数量和合约乘数100进行加总。这样得到的是组合层面的风险暴露，而不是单个合约的孤立敏感度。",
    )
    add_bullet(doc, f"组合Delta={num(total['portfolio_delta'], 2)}：标的每上涨1美元，组合理论价值约变化该数值美元。")
    add_bullet(doc, f"组合Gamma={num(total['portfolio_gamma'], 4)}：Delta随标的价格变化的速度，价差策略通常低于单腿期权。")
    add_bullet(doc, f"组合Vega={num(total['portfolio_vega'], 2)}：IV每变化1个百分点，组合价值约变化该数值美元。")
    add_bullet(doc, f"组合Theta={num(total['portfolio_theta'], 2)}：每日时间价值影响，本策略为 {num(total['portfolio_theta'], 2)}。")
    if (OUTPUT_DIR / "portfolio_greeks.png").exists():
        doc.add_picture(str(OUTPUT_DIR / "portfolio_greeks.png"), width=Inches(5.5))
        cap = doc.add_paragraph("图2  组合策略希腊字母")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "四、模块B：策略盈亏分析")
    add_para(
        doc,
        f"新增模块B包含两层盈亏分析：第一层是到期日静态盈亏，第二层是7天后的标的价格和IV情景重估。当前组合初始净成本为 {money(pnl['initial_value'])} 美元；在网格范围内最大盈利约 {money(pnl['max_profit_in_grid'])} 美元，最大亏损约 {money(pnl['max_loss_in_grid'])} 美元，盈亏平衡点约 {', '.join(num(x) for x in pnl['breakevens_in_grid'])}。",
    )
    if (OUTPUT_DIR / "strategy_payoff.png").exists():
        doc.add_picture(str(OUTPUT_DIR / "strategy_payoff.png"), width=Inches(5.7))
        cap = doc.add_paragraph("图3  到期盈亏曲线")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if (OUTPUT_DIR / "scenario_pnl_heatmap.png").exists():
        doc.add_picture(str(OUTPUT_DIR / "scenario_pnl_heatmap.png"), width=Inches(5.2))
        cap = doc.add_paragraph("图4  7天后价格/IV情景盈亏")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    best = scenario.loc[scenario["pnl_after_7d"].idxmax()]
    worst = scenario.loc[scenario["pnl_after_7d"].idxmin()]
    add_bullet(doc, f"最佳情景：标的变动 {pct(best['spot_move'])}，IV变动 {pct(best['iv_shift'])}，7天后P&L约 {money(best['pnl_after_7d'])} 美元。")
    add_bullet(doc, f"最差情景：标的变动 {pct(worst['spot_move'])}，IV变动 {pct(worst['iv_shift'])}，7天后P&L约 {money(worst['pnl_after_7d'])} 美元。")

    if multi_horizon is not None and not multi_horizon.empty:
        add_heading(doc, "五、模块B扩展1：多时间断面情景")
        add_para(
            doc,
            "在保持隐含波动率不变的前提下，把策略价值在持仓1、7、14、21天后的5档价格点上重新定价，可以直观看到Theta消耗与价格弹性的相互作用。",
        )
        horizons = sorted(multi_horizon["horizon_days"].unique().tolist())
        spot_moves = sorted(multi_horizon["spot_move"].unique().tolist())
        mh_table = doc.add_table(rows=1, cols=1 + len(spot_moves))
        mh_table.rows[0].cells[0].text = "持仓天数 \\ 标的变动"
        for j, m in enumerate(spot_moves):
            mh_table.rows[0].cells[j + 1].text = pct(m)
        for h in horizons:
            row = mh_table.add_row().cells
            row[0].text = f"{int(h)} 天"
            for j, m in enumerate(spot_moves):
                cell = multi_horizon[(multi_horizon["horizon_days"] == h) & (multi_horizon["spot_move"] == m)]
                row[j + 1].text = money(float(cell["pnl"].iloc[0]))
        style_table(mh_table)
        if (OUTPUT_DIR / "multi_horizon_pnl.png").exists():
            doc.add_picture(str(OUTPUT_DIR / "multi_horizon_pnl.png"), width=Inches(5.7))
            cap = doc.add_paragraph("图5  多时间断面盈亏曲线")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        center_now = multi_horizon[multi_horizon["spot_move"] == 0.0].sort_values("horizon_days")
        if len(center_now) >= 2:
            theta_drift = float(center_now["pnl"].iloc[-1] - center_now["pnl"].iloc[0])
            add_bullet(doc, f"标的不动且IV不变时，从持仓1天到{int(center_now['horizon_days'].iloc[-1])}天的P&L变化约为 {money(theta_drift)} 美元，体现Theta累计影响。")

    if mc_stats is not None and not mc_stats.empty:
        add_heading(doc, "六、模块B扩展2：Monte Carlo盈亏与风险")
        rn = mc_stats[mc_stats["scenario"] == "risk_neutral"].iloc[0]
        rw = mc_stats[mc_stats["scenario"] == "real_world"].iloc[0]
        add_para(
            doc,
            "为了量化策略到期盈亏的概率分布，对标的资产做对数正态Monte Carlo模拟（共20000条路径）。同时给出风险中性（drift=r=4%, sigma=ATM IV）与真实世界（drift=封顶后的近20日年化收益, sigma=RV60）两种漂移假设：前者代表市场公允定价隐含的成功率，后者代表近期价格趋势延续下的成功率。",
        )
        mc_table = doc.add_table(rows=1, cols=8)
        for idx, h in enumerate([
            "情景",
            "drift",
            "sigma",
            "POP",
            "E[P&L]",
            "中位数",
            "VaR95",
            "CVaR95",
        ]):
            mc_table.rows[0].cells[idx].text = h
        for label, row in [("风险中性", rn), ("真实世界", rw)]:
            cells = mc_table.add_row().cells
            cells[0].text = label
            cells[1].text = pct(row["annual_drift_used"])
            cells[2].text = pct(row["annual_sigma_used"])
            cells[3].text = pct(row["prob_of_profit"])
            cells[4].text = money(row["expected_pnl"])
            cells[5].text = money(row["median_pnl"])
            cells[6].text = money(row["var_95"])
            cells[7].text = money(row["cvar_95"])
        style_table(mc_table)
        if (OUTPUT_DIR / "monte_carlo_pnl.png").exists():
            doc.add_picture(str(OUTPUT_DIR / "monte_carlo_pnl.png"), width=Inches(5.9))
            cap = doc.add_paragraph("图6  风险中性Monte Carlo盈亏分布")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_bullet(doc, f"风险中性POP={pct(rn['prob_of_profit'])}，期望盈亏≈{money(rn['expected_pnl'])}美元，体现市场对该结构的公允定价。")
        add_bullet(doc, f"真实世界POP={pct(rw['prob_of_profit'])}，期望盈亏≈{money(rw['expected_pnl'])}美元，反映近期上涨趋势带来的方向性收益。")
        add_bullet(doc, f"VaR95和CVaR95均等于权利金净支出 {money(rn['var_95'])}美元，符合牛市价差最大亏损被锁定为权利金的金融常识。")

    add_heading(doc, "七、结论")
    add_para(
        doc,
        "本次扩展使 wisecoin_option 类程序从单纯下载期权链和给出方向观点，升级为完整的策略评估流程：观点形成后能自动选组合，组合层面能看到Delta/Vega/Theta等风险暴露，盈亏模块从到期静态曲线、IV情景重估扩展到多时间断面情景与Monte Carlo POP/VaR/CVaR。这样既能定量回答\"现在赚钱概率多大、最坏可能亏多少\"，也能区分\"市场公允定价隐含的胜率\"与\"个人观点下的胜率\"。结论仅用于课程作业演示，真实交易还需考虑滑点、手续费、成交量、提前行权和风险预算。",
    )

    add_heading(doc, "八、附件清单")
    add_bullet(doc, "code/wisecoin_option_analysis.py：数据下载、观点形成、组合Greeks、到期盈亏、IV情景、多时间断面、Monte Carlo盈亏与风险。")
    add_bullet(doc, "code/generate_wisecoin_option_homework_report.py：生成Word/PDF和提交包。")
    add_bullet(doc, "data/：SPY历史行情与期权链CSV。")
    add_bullet(doc, "outputs/：策略腿、组合Greeks、盈亏表、Monte Carlo路径与统计、多时间断面表、观点摘要和图形（共6张）。")

    doc.sections[0].footer.paragraphs[0].text = f"{SUBMISSION_NAME} | wisecoin_option程序分析"
    doc.save(REPORT_DOCX)


def register_fonts() -> tuple[str, str]:
    regular = r"C:\Windows\Fonts\msyh.ttc"
    bold = r"C:\Windows\Fonts\msyhbd.ttc"
    pdfmetrics.registerFont(TTFont("MSYH", regular))
    pdfmetrics.registerFont(TTFont("MSYH-Bold", bold))
    return "MSYH", "MSYH-Bold"


def p(text: str, style) -> Paragraph:
    return Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style)


def build_pdf() -> None:
    data = load_outputs()
    view = data["view"]
    legs = data["legs"]
    greeks = data["greeks"]
    scenario = data["scenario"]
    mc_stats = data["mc_stats"]
    multi_horizon = data["multi_horizon"]
    total = greeks[greeks["contract_symbol"] == "TOTAL"].iloc[0]
    pnl = view["pnl_summary"]
    font, bold = register_fonts()

    doc = SimpleDocTemplate(str(REPORT_PDF), pagesize=A4, rightMargin=1.5 * cm, leftMargin=1.5 * cm, topMargin=1.4 * cm, bottomMargin=1.4 * cm)
    styles = getSampleStyleSheet()
    title = ParagraphStyle("title", parent=styles["Title"], fontName=bold, fontSize=17, leading=23, alignment=TA_CENTER, textColor=colors.HexColor("#1F4E79"))
    h = ParagraphStyle("h", parent=styles["Heading1"], fontName=bold, fontSize=12.5, leading=17, spaceBefore=9, spaceAfter=5, textColor=colors.HexColor("#1F4E79"))
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName=font, fontSize=9.1, leading=13.5, firstLineIndent=18, spaceAfter=4)
    cap = ParagraphStyle("cap", parent=styles["BodyText"], fontName=font, fontSize=8.2, leading=11, alignment=TA_CENTER)

    story = [
        Paragraph("wisecoin_option程序分析、组合希腊字母与策略盈亏", title),
        p(f"提交文件名：{SUBMISSION_NAME}    数据样本：{view['ticker']} Yahoo期权链", body),
        Paragraph("一、程序审查", h),
        p(f"本作业复现wisecoin_option类程序的数据下载与观点形成流程：使用yfinance下载{view['ticker']}近一年价格和DTE接近45天的期权链，按中间价清洗报价，并根据MA20/MA60、20日收益、RSI和IV/RV形成观点。样本日{view['as_of']}，到期日{view['expiry']}，方向观点为{view['directional_view_cn']}，波动率观点为{view['volatility_view']}。", body),
        Paragraph("二、策略与组合希腊字母", h),
        p(f"程序选择{view['strategy_name']}。组合Delta={num(total['portfolio_delta'], 2)}，Gamma={num(total['portfolio_gamma'], 4)}，Vega={num(total['portfolio_vega'], 2)}，Theta={num(total['portfolio_theta'], 2)}，Rho={num(total['portfolio_rho'], 2)}。", body),
    ]
    legs_table = [["合约", "类型", "K", "方向", "中间价", "IV"]]
    for _, row in legs.iterrows():
        legs_table.append([
            str(row["contract_symbol"]),
            str(row["option_type"]),
            num(row["strike"], 0),
            "买入" if int(row["quantity"]) > 0 else "卖出",
            money(row["mid"]),
            pct(row["implied_volatility"]),
        ])
    table = Table(legs_table, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (-1, -1), font, 6.8),
                ("FONT", (0, 0), (-1, 0), bold, 7),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ]
        )
    )
    story.extend([table, Spacer(1, 5)])
    for img, caption in [
        ("underlying_view.png", "图1 标的价格与趋势指标"),
        ("portfolio_greeks.png", "图2 组合希腊字母"),
        ("strategy_payoff.png", "图3 到期盈亏曲线"),
        ("scenario_pnl_heatmap.png", "图4 7天后情景盈亏"),
    ]:
        path = OUTPUT_DIR / img
        if path.exists():
            story.extend([Image(str(path), width=14.2 * cm, height=6.2 * cm), Paragraph(caption, cap), Spacer(1, 4)])

    best = scenario.loc[scenario["pnl_after_7d"].idxmax()]
    worst = scenario.loc[scenario["pnl_after_7d"].idxmin()]
    story.extend(
        [
            Paragraph("三、盈亏结论（到期与7天IV情景）", h),
            p(f"初始净成本为{money(pnl['initial_value'])}美元，网格内最大盈利约{money(pnl['max_profit_in_grid'])}美元，最大亏损约{money(pnl['max_loss_in_grid'])}美元，盈亏平衡点约{', '.join(num(x) for x in pnl['breakevens_in_grid'])}。最佳7天情景P&L为{money(best['pnl_after_7d'])}美元，最差7天情景P&L为{money(worst['pnl_after_7d'])}美元。", body),
        ]
    )

    if multi_horizon is not None and not multi_horizon.empty:
        horizons = sorted(multi_horizon["horizon_days"].unique().tolist())
        spot_moves = sorted(multi_horizon["spot_move"].unique().tolist())
        mh_table_data = [["持仓天数 \\ 标的"] + [pct(m) for m in spot_moves]]
        for hh in horizons:
            row = [f"{int(hh)}天"]
            for m in spot_moves:
                cell = multi_horizon[(multi_horizon["horizon_days"] == hh) & (multi_horizon["spot_move"] == m)]
                row.append(money(float(cell["pnl"].iloc[0])))
            mh_table_data.append(row)
        mh_table = Table(mh_table_data, repeatRows=1)
        mh_table.setStyle(
            TableStyle(
                [
                    ("FONT", (0, 0), (-1, -1), font, 7.0),
                    ("FONT", (0, 0), (-1, 0), bold, 7.2),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            )
        )
        story.extend(
            [
                Paragraph("四、模块B扩展1：多时间断面情景", h),
                p("保持隐含波动率不变，将策略价值在持仓1、7、14、21天后的5档价格点上重新定价，便于直观比较Theta消耗与价格弹性的相互作用。", body),
                mh_table,
                Spacer(1, 4),
            ]
        )
        path = OUTPUT_DIR / "multi_horizon_pnl.png"
        if path.exists():
            story.extend([Image(str(path), width=14.2 * cm, height=6.2 * cm), Paragraph("图5 多时间断面盈亏曲线", cap), Spacer(1, 4)])

    if mc_stats is not None and not mc_stats.empty:
        rn = mc_stats[mc_stats["scenario"] == "risk_neutral"].iloc[0]
        rw = mc_stats[mc_stats["scenario"] == "real_world"].iloc[0]
        mc_table_data = [["情景", "drift", "sigma", "POP", "E[P&L]", "中位数", "VaR95", "CVaR95"]]
        for label, row in [("风险中性", rn), ("真实世界", rw)]:
            mc_table_data.append(
                [
                    label,
                    pct(row["annual_drift_used"]),
                    pct(row["annual_sigma_used"]),
                    pct(row["prob_of_profit"]),
                    money(row["expected_pnl"]),
                    money(row["median_pnl"]),
                    money(row["var_95"]),
                    money(row["cvar_95"]),
                ]
            )
        mc_table = Table(mc_table_data, repeatRows=1)
        mc_table.setStyle(
            TableStyle(
                [
                    ("FONT", (0, 0), (-1, -1), font, 7.0),
                    ("FONT", (0, 0), (-1, 0), bold, 7.2),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            )
        )
        story.extend(
            [
                Paragraph("五、模块B扩展2：Monte Carlo盈亏与风险", h),
                p("对标的资产做对数正态Monte Carlo模拟（20000条路径）。风险中性假设drift=r=4%、sigma=ATM IV，反映市场公允定价隐含的成功率；真实世界假设drift=封顶后近20日年化收益、sigma=RV60，反映近期价格趋势延续下的成功率。", body),
                mc_table,
                Spacer(1, 4),
                p(f"风险中性POP={pct(rn['prob_of_profit'])}，期望盈亏≈{money(rn['expected_pnl'])}美元；真实世界POP={pct(rw['prob_of_profit'])}，期望盈亏≈{money(rw['expected_pnl'])}美元。VaR95与CVaR95均等于权利金净支出，符合牛市价差最大亏损被锁定的金融常识。", body),
            ]
        )
        path = OUTPUT_DIR / "monte_carlo_pnl.png"
        if path.exists():
            story.extend([Image(str(path), width=14.2 * cm, height=6.2 * cm), Paragraph("图6 风险中性Monte Carlo盈亏分布", cap), Spacer(1, 4)])

    story.extend(
        [
            Paragraph("六、结论", h),
            p("本作业把wisecoin_option类程序扩展为完整的期权策略研究流程：下载数据、形成观点、自动选组合、组合层面Greeks、到期与IV情景盈亏、多时间断面、Monte Carlo POP/VaR/CVaR。既能定量回答\"现在赚钱概率多大、最坏可能亏多少\"，也能区分\"市场公允定价隐含的胜率\"与\"个人观点下的胜率\"。结果用于课程展示，真实交易需加入成交量、滑点、手续费和风险限额。", body),
        ]
    )
    doc.build(story)


def make_zip() -> Path:
    zip_path = ROOT / "作业提交包" / f"{SUBMISSION_NAME}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in PACKAGE_DIR.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(PACKAGE_PARENT))
    return zip_path


def main() -> None:
    ensure_package_dir()
    copy_artifacts()
    write_readme()
    build_docx()
    build_pdf()
    zip_path = make_zip()
    print(f"DOCX: {REPORT_DOCX}")
    print(f"PDF: {REPORT_PDF}")
    print(f"ZIP: {zip_path}")


if __name__ == "__main__":
    main()
