"""
generate_week10_project_report.py
=================================

Generate the Week-10 Project submission package (DOCX + PDF + ZIP).

The project is a historical option-data reconstruction + daily-bar simulated
trading strategy on SPY (2024-2026 window).  This generator turns the
backtest's outputs into a polished, publication-style report.
"""

from __future__ import annotations

import json
import shutil
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_NAME = "23300180062_任昊来_第十周Project"
PROJECT_OUT = ROOT / "research/outputs/week10/project"
DATA_DIR = ROOT / "data/week10"
PACKAGE_PARENT = ROOT / "作业提交包" / "_staging_week10_proj"
PACKAGE_DIR = PACKAGE_PARENT / SUBMISSION_NAME
DOCX_PATH = PACKAGE_DIR / f"{SUBMISSION_NAME}.docx"
PDF_PATH = PACKAGE_DIR / f"{SUBMISSION_NAME}.pdf"


def num(x, d=2):
    if pd.isna(x):
        return ""
    return f"{float(x):.{d}f}"


def money(x):
    if pd.isna(x):
        return ""
    return f"{float(x):,.2f}"


def pct(x):
    if pd.isna(x):
        return ""
    return f"{float(x):.2%}"


def ensure_pkg() -> None:
    if PACKAGE_DIR.exists():
        shutil.rmtree(PACKAGE_DIR)
    for sub in ["code", "data", "outputs"]:
        (PACKAGE_DIR / sub).mkdir(parents=True, exist_ok=True)


def copy_artifacts() -> None:
    code_src = [
        "research/week10_project_backtest.py",
        "research/generate_week10_project_report.py",
    ]
    for rel in code_src:
        src = ROOT / rel
        if src.exists():
            shutil.copy2(src, PACKAGE_DIR / "code" / src.name)
    for path in DATA_DIR.glob("*_history*"):
        if path.is_file():
            shutil.copy2(path, PACKAGE_DIR / "data" / path.name)
    for path in PROJECT_OUT.glob("*"):
        if path.is_file():
            shutil.copy2(path, PACKAGE_DIR / "outputs" / path.name)


def write_readme(stats: dict) -> None:
    (PACKAGE_DIR / "README.txt").write_text(
        "\n".join(
            [
                "第十周 Project：历史期权数据再现 + 日线仿真交易策略回测",
                "",
                f"提交名称：{SUBMISSION_NAME}",
                "",
                "项目目标：",
                "1. 在缺乏免费完整历史期权链数据的情况下，",
                "   用 Black-Scholes 模型 + 实现波动率拟合 IV 曲面，",
                "   从标的资产日线价格再现 2024-2026 两年的虚拟期权链。",
                "2. 设计并回测一个仿真交易策略：滚动卖出 30 天到期、0.20 delta 的宽跨（Short Strangle），",
                "   含开仓阈值、止盈 (50%)、止损 (200%)、管理期 (DTE=2) 退出规则。",
                "3. 多维度量化损益和风险：净值曲线、回撤、交易级 P&L 分布、",
                "   日度 Greeks 暴露、与买入持有的对比。",
                "",
                "主要文件：",
                f"- {SUBMISSION_NAME}.pdf：PDF 报告",
                f"- {SUBMISSION_NAME}.docx：Word 报告",
                "- code/week10_project_backtest.py：完整回测引擎（数据 → IV 曲面 → 策略 → 风控 → 统计）",
                "- code/generate_week10_project_report.py：本报告生成器",
                "- data/：标的资产历史与 RV 加工后的数据",
                "- outputs/：净值、交易、统计表、所有图形",
                "",
                "关键统计：",
                f"- 回测窗口：{stats['start_date']} → {stats['end_date']} ({stats['n_days']} 天)",
                f"- 总交易次数：{stats['n_trades']}，胜率：{stats['win_rate']:.1%}",
                f"- 策略 CAGR：{stats['cagr']:.2%} vs 买入持有 CAGR：{stats['buyhold_cagr']:.2%}",
                f"- 最大回撤：{stats['max_drawdown']:.2%}，Sharpe：{stats['sharpe_ratio']:.2f}",
                "",
                "复现命令：",
                "python research/week10_project_backtest.py",
                "python research/generate_week10_project_report.py",
            ]
        ),
        encoding="utf-8",
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.5)
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, "1F4E79")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(18)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.3)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.8)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.6)
    r.font.color.rgb = RGBColor(60, 60, 60)


def build_docx() -> None:
    stats = json.loads((PROJECT_OUT / "stats.json").read_text(encoding="utf-8"))
    trades = pd.read_csv(PROJECT_OUT / "trades.csv")

    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(1.9)
    s.right_margin = Cm(1.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("第十周 Project：历史期权数据再现 + 日线仿真交易策略回测")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    meta = doc.add_paragraph(
        f"提交名称：{SUBMISSION_NAME}    标的：{stats['ticker']}    数据源：{stats['data_source']}    "
        f"回测窗口：{stats['start_date']} → {stats['end_date']}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)

    # Abstract
    add_heading(doc, "摘要", level=1)
    add_para(
        doc,
        f"本项目把 SPY ETF 在过去两年（{stats['start_date']} 至 {stats['end_date']}）的日线价格序列还原为虚拟的历史期权链，"
        "并在其上构建并回测了一个仿真交易策略：每周五滚动卖出 30 天到期、±0.20 delta 的宽跨。"
        "策略带止盈（50% 最大盈利）、止损（200% 初始权利金）、管理期（DTE=2）三重风控。"
        f"两年时间共执行 {stats['n_trades']} 笔交易，胜率 {stats['win_rate']:.1%}，"
        f"年化收益 {stats['cagr']:.2%}（vs 买入持有 {stats['buyhold_cagr']:.2%}），"
        f"最大回撤 {stats['max_drawdown']:.2%}，Sharpe {stats['sharpe_ratio']:.2f}。"
        "尽管胜率高，但单笔大亏几乎抹平所有盈利，体现了卖方策略的尾部风险特征。",
    )

    # I. Motivation
    add_heading(doc, "一、问题背景与项目目标", level=1)
    add_para(
        doc,
        "完整历史期权 tick 级数据 (CBOE DataShop、Polygon Options) 价格不菲且无免费渠道；学术研究中"
        "通常采用「数据再现」方法：用标的日线价格 + 实现波动率 + 一组合理假设的 IV 曲面，"
        "用 Black-Scholes 反推每日的虚拟期权链。这种方法可以让我们在没有真实 tick 数据的情况下，研究："
    )
    add_bullet(doc, "卖方期权策略（Short Strangle / Iron Condor 等）的长期表现")
    add_bullet(doc, "策略风控参数（止盈/止损/管理期）对收益的边际贡献")
    add_bullet(doc, "在不同市场状态下（趋势 vs 震荡，低波 vs 高波）的损益分布")
    add_bullet(doc, "组合层面的日度 Greeks 暴露与对冲需求")

    # II. Data Reconstruction
    add_heading(doc, "二、历史期权数据再现方法", level=1)
    add_heading(doc, "2.1 数据流程", level=2)
    add_para(
        doc,
        "数据再现分四步：(1) yfinance 拉取 SPY 三年日线；"
        f"(2) 计算 20/60/120 日年化实现波动率 RV；(3) 在每个交易日生成 5 个到期日（7/14/30/45/60 DTE）的虚拟期权链；"
        f"(4) 用 BS 模型 + 实现波动率拟合的 IV 曲面定价每个合约。",
    )
    add_heading(doc, "2.2 IV 曲面假设", level=2)
    add_para(doc, "IV 曲面公式（在 week10_project_backtest.py 中实现）：")
    add_code_block(
        doc,
        "IV(t, K, T) = max( 0.05,\n"
        "                  RV_20(t) * IV_RV_RATIO\n"
        "                  - PUT_SKEW * ln(K / F_t) / sqrt(T) )",
    )
    add_bullet(doc, f"IV_RV_RATIO = {stats['iv_rv_ratio']:.2f}：观察到的隐含波动率对实现波动率的平均溢价。")
    add_bullet(doc, "PUT_SKEW = 0.06：股票指数典型的看跌偏度（put 的 IV 在远端高于 call）。")
    add_bullet(
        doc,
        "下层 0.05 截断：避免 RV 接近零时（极端低波动期）生成不合理的接近零的 IV。",
    )
    if (PROJECT_OUT / "iv_rv_history.png").exists():
        doc.add_picture(str(PROJECT_OUT / "iv_rv_history.png"), width=Inches(6.4))
        cap = doc.add_paragraph(f"图1  回测期内 SPY 的 RV_20 与拟合的 IV（IV = RV_20 × {stats['iv_rv_ratio']:.2f}）")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # III. Strategy
    add_heading(doc, "三、仿真交易策略与回测引擎", level=1)
    add_heading(doc, "3.1 策略规则", level=2)
    add_bullet(doc, f"开仓：每个交易周的周五，若当前空仓，则按目标 delta = ±{stats['entry_delta']:.2f} 选 30 天到期的看涨与看跌期权，构建宽跨卖出仓位。")
    add_bullet(doc, f"仓位规模：每笔交易锁定 $50,000 现金作担保，按行权价 × 100 计算合约数。")
    add_bullet(doc, f"止盈：当 MTM 达到初始权利金的 {stats['take_profit']:.0%}（即 50% 最大盈利）时平仓。")
    add_bullet(doc, f"止损：当 MTM 损失达到初始权利金的 {stats['stop_loss']:.0%} 时平仓，防止尾部失控。")
    add_bullet(doc, f"管理期：当 DTE ≤ {stats['managed_dte']} 时强制平仓，避免短期 gamma 风险与 pin risk。")
    add_bullet(doc, "成本：每条腿 $0.65 佣金 × 开平合计 4 次；额外 1% 中点价的 bid-ask 滑点。")
    add_heading(doc, "3.2 引擎核心循环", level=2)
    add_para(
        doc,
        "回测引擎按日推进，每日依次完成：1) 用当日 RV 重建 IV 曲面与当日虚拟期权链；"
        "2) 若持仓在线则重定价、检查风控三阈值；3) 若空仓且为周五则按规则开仓；"
        "4) 记录现金、MTM、组合 Greeks 与净值。",
    )

    # IV. Results
    add_heading(doc, "四、回测结果", level=1)
    add_heading(doc, "4.1 净值曲线与回撤", level=2)
    if (PROJECT_OUT / "equity_curve.png").exists():
        doc.add_picture(str(PROJECT_OUT / "equity_curve.png"), width=Inches(6.4))
        cap = doc.add_paragraph("图2  策略净值 vs SPY 买入持有")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if (PROJECT_OUT / "drawdown.png").exists():
        doc.add_picture(str(PROJECT_OUT / "drawdown.png"), width=Inches(6.4))
        cap = doc.add_paragraph("图3  策略最大回撤序列")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "4.2 关键统计指标", level=2)
    stats_tab = doc.add_table(rows=1, cols=2)
    stats_tab.rows[0].cells[0].text = "指标"
    stats_tab.rows[0].cells[1].text = "数值"
    rows = [
        ("初始资金", money(stats["start_equity"])),
        ("结束净值", money(stats["end_equity"])),
        ("总收益率", pct(stats["total_return"])),
        ("年化收益率 (策略)", pct(stats["cagr"])),
        ("年化收益率 (买入持有)", pct(stats["buyhold_cagr"])),
        ("Sharpe 比率", num(stats["sharpe_ratio"], 2)),
        ("Sortino 比率", num(stats["sortino_ratio"], 2)),
        ("最大回撤", pct(stats["max_drawdown"])),
        ("交易次数", str(stats["n_trades"])),
        ("胜率", pct(stats["win_rate"])),
        ("平均盈利", money(stats["avg_win"])),
        ("平均亏损", money(stats["avg_loss"])),
        ("期望 P&L / 笔", money(stats["expectancy"])),
        ("Profit Factor", num(stats["profit_factor"], 2)),
        ("最大单笔盈利", money(stats["max_trade_pnl"])),
        ("最大单笔亏损", money(stats["min_trade_pnl"])),
    ]
    for k, v in rows:
        c = stats_tab.add_row().cells
        c[0].text = k
        c[1].text = v
    style_table(stats_tab)

    add_heading(doc, "4.3 交易级 P&L 分布", level=2)
    if (PROJECT_OUT / "trade_pnl.png").exists():
        doc.add_picture(str(PROJECT_OUT / "trade_pnl.png"), width=Inches(6.4))
        cap = doc.add_paragraph("图4  左：交易 P&L 直方图  右：累计交易 P&L")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "4.4 组合 Greeks 暴露随时间变化", level=2)
    if (PROJECT_OUT / "greek_exposure.png").exists():
        doc.add_picture(str(PROJECT_OUT / "greek_exposure.png"), width=Inches(6.4))
        cap = doc.add_paragraph("图5  日度 Delta / Gamma / Vega / Theta 暴露（持仓为空时为 0）")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # V. Trade examples
    add_heading(doc, "五、典型交易样本", level=1)
    if not trades.empty:
        # show 6 representative trades: 2 best wins, 2 worst losses, 2 mid
        top = trades.nlargest(2, "pnl")
        bottom = trades.nsmallest(2, "pnl")
        median = trades.iloc[[len(trades) // 2 - 1, len(trades) // 2]] if len(trades) >= 4 else trades.head(2)
        sample = pd.concat([top, bottom, median]).drop_duplicates(subset=["open_date", "expiry"]).sort_values("open_date")
        t_tab = doc.add_table(rows=1, cols=8)
        for i, h in enumerate([
            "开仓",
            "平仓",
            "天数",
            "Put K",
            "Call K",
            "开仓权利金",
            "P&L",
            "退出原因",
        ]):
            t_tab.rows[0].cells[i].text = h
        for _, r in sample.iterrows():
            c = t_tab.add_row().cells
            c[0].text = str(r["open_date"])
            c[1].text = str(r["close_date"])
            c[2].text = str(int(r["days_held"]))
            c[3].text = num(r["put_strike"], 1)
            c[4].text = num(r["call_strike"], 1)
            c[5].text = money(r["open_credit"])
            c[6].text = money(r["pnl"])
            c[7].text = str(r["exit_reason"])
        style_table(t_tab)

    # VI. Discussion
    add_heading(doc, "六、讨论与改进方向", level=1)
    add_heading(doc, "6.1 卖方策略的「拣硬币」困境", level=2)
    add_para(
        doc,
        f"本回测期内策略胜率达到 {stats['win_rate']:.1%}，但 Profit Factor 仅 {stats['profit_factor']:.2f}。"
        f"平均盈利 {money(stats['avg_win'])} vs 平均亏损 {money(stats['avg_loss'])}，亏损是盈利的 "
        f"{abs(stats['avg_loss'] / stats['avg_win']):.2f} 倍。"
        "这就是卖方策略经典的「在压路机前拣硬币」效应：长期累积小盈利、但偶发的大行情可以一次性吃掉几个月的利润。"
        "相比之下买入持有 SPY 在同期年化 "
        f"{stats['buyhold_cagr']:.2%}，体现了股票指数长期向上的阿尔法是卖期权策略难以战胜的基准。",
    )
    add_heading(doc, "6.2 IV 曲面假设的局限", level=2)
    add_para(
        doc,
        f"本仿真使用 IV = RV_20 × {stats['iv_rv_ratio']:.2f} 作为基础 IV 估计，"
        "这是历史平均的 IV-RV 溢价，但在不同市场状态下溢价波动很大："
        "VIX 短期飙升时实际 IV 可远超此倍数（行情亏损时空头权利金贬值更慢，本仿真低估了对冲难度）；"
        "在低波动稳定期 IV-RV 溢价又可能压缩到 0.95 左右（高估了实际权利金）。"
        "改进方向：将 IV_RV_RATIO 改为状态依赖的回归模型（参数依赖 VIX/SKEW/ trend）。",
    )
    add_heading(doc, "6.3 进一步研究方向", level=2)
    add_bullet(doc, "对卖方策略加入 Delta 中性对冲（每日用标的期货回填 |Δ| > 阈值的暴露）。")
    add_bullet(doc, "用 IV 高分位过滤入场：只在 IV percentile > 50 时开仓，避开低波动「无奖励」情境。")
    add_bullet(doc, "Iron Condor 替代 Short Strangle：以一定盈利空间换取尾部锁定。")
    add_bullet(doc, "参数最优化：在 entry_delta ∈ {0.10, 0.15, 0.20, 0.30}、take_profit ∈ {0.25, 0.50, 0.75} 的网格上做敏感性分析。")
    add_bullet(doc, "升级到真实数据：接入 Polygon Options 或 CBOE DataShop 后比较「合成 IV」与真实 IV 的差异。")

    add_heading(doc, "七、结论", level=1)
    add_para(
        doc,
        "本项目完整闭环地实现了「数据再现 → 期权链合成 → 策略仿真 → 风控执行 → 风险/收益归因」"
        "整套量化研究流程。"
        f"在 SPY 2024-2026 样本中，胜率 {stats['win_rate']:.1%} 的滚动短宽跨策略最终"
        f"取得 {stats['total_return']:.2%} 总收益，远低于 SPY 买入持有 {stats['buyhold_total_return']:.2%}。"
        "这一结果不是策略失败的证据，而是揭示了无对冲、无 IV-percentile 过滤的纯卖方策略"
        "在牛市中的系统性劣势——尾部一次「黑天鹅」可以抹平多个月的小赚。"
        "后续研究将重点放在 Delta 对冲、状态依赖入场与 IV 曲面动态拟合上。",
    )

    add_heading(doc, "八、附件清单", level=1)
    add_bullet(doc, "code/week10_project_backtest.py：完整回测引擎（数据→IV曲面→策略→风控→统计）")
    add_bullet(doc, "code/generate_week10_project_report.py：本报告生成器")
    add_bullet(doc, "data/spy_history_backtest.csv, spy_history_with_rv.csv：标的资产历史与 RV 加工后数据")
    add_bullet(doc, "outputs/equity_curve.csv, trades.csv, stats.csv, stats.json")
    add_bullet(doc, "outputs/equity_curve.png, drawdown.png, trade_pnl.png, greek_exposure.png, iv_rv_history.png")

    doc.sections[0].footer.paragraphs[0].text = f"{SUBMISSION_NAME} | 第十周 Project"
    doc.save(DOCX_PATH)


def register_fonts() -> tuple[str, str]:
    pdfmetrics.registerFont(TTFont("MSYH", r"C:\Windows\Fonts\msyh.ttc"))
    pdfmetrics.registerFont(TTFont("MSYH-Bold", r"C:\Windows\Fonts\msyhbd.ttc"))
    return "MSYH", "MSYH-Bold"


def p_safe(text, style):
    return Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style)


def build_pdf() -> None:
    stats = json.loads((PROJECT_OUT / "stats.json").read_text(encoding="utf-8"))
    trades = pd.read_csv(PROJECT_OUT / "trades.csv")
    font, bold = register_fonts()

    doc = SimpleDocTemplate(
        str(PDF_PATH), pagesize=A4, rightMargin=1.4 * cm, leftMargin=1.4 * cm, topMargin=1.3 * cm, bottomMargin=1.3 * cm
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("title", parent=styles["Title"], fontName=bold, fontSize=15.5, leading=21, alignment=TA_CENTER, textColor=colors.HexColor("#1F4E79"))
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName=bold, fontSize=12.5, leading=17, spaceBefore=10, spaceAfter=5, textColor=colors.HexColor("#1F4E79"))
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName=bold, fontSize=10.5, leading=14, spaceBefore=6, spaceAfter=3, textColor=colors.HexColor("#1F4E79"))
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName=font, fontSize=8.9, leading=12.7, firstLineIndent=15, spaceAfter=3)
    bullet = ParagraphStyle("bul", parent=styles["BodyText"], fontName=font, fontSize=8.5, leading=12.0, leftIndent=15, spaceAfter=2)
    cap = ParagraphStyle("cap", parent=styles["BodyText"], fontName=font, fontSize=8.0, leading=10.5, alignment=TA_CENTER)
    code = ParagraphStyle("code", parent=styles["BodyText"], fontName="Courier", fontSize=8.4, leading=11.2, leftIndent=20, spaceAfter=4)
    abstract = ParagraphStyle("abs", parent=styles["BodyText"], fontName=font, fontSize=9.0, leading=13.2, spaceAfter=4, textColor=colors.HexColor("#333333"))

    def std_table(data: list[list[str]]) -> Table:
        t = Table(data, repeatRows=1)
        t.setStyle(
            TableStyle(
                [
                    ("FONT", (0, 0), (-1, -1), font, 7.2),
                    ("FONT", (0, 0), (-1, 0), bold, 7.4),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            )
        )
        return t

    def add_image(path: Path, caption: str, w_cm: float, h_cm: float, story: list) -> None:
        if path.exists():
            story.append(Image(str(path), width=w_cm * cm, height=h_cm * cm))
            story.append(Paragraph(caption, cap))
            story.append(Spacer(1, 3))

    story = [
        Paragraph("第十周 Project：历史期权数据再现 + 日线仿真交易策略回测", title),
        p_safe(
            f"提交名称：{SUBMISSION_NAME}    标的：{stats['ticker']}    数据源：{stats['data_source']}    回测窗口：{stats['start_date']} → {stats['end_date']}",
            body,
        ),
        Paragraph("摘要", h1),
        p_safe(
            f"本项目把 SPY 在 {stats['start_date']} 至 {stats['end_date']} 的日线价格序列还原为虚拟的历史期权链，"
            "并在其上构建并回测了一个仿真交易策略：每周五滚动卖出 30 天到期、±0.20 delta 的宽跨。"
            "策略带止盈（50%）、止损（200%）、管理期（DTE=2）三重风控。"
            f"共执行 {stats['n_trades']} 笔交易，胜率 {stats['win_rate']:.1%}；"
            f"年化收益 {stats['cagr']:.2%}（vs 买入持有 {stats['buyhold_cagr']:.2%}），"
            f"最大回撤 {stats['max_drawdown']:.2%}，Sharpe {stats['sharpe_ratio']:.2f}。"
            "结果展示了卖方策略「胜率高但偶发大亏抹平利润」的经典特征。",
            abstract,
        ),
    ]

    # Section I
    story.extend(
        [
            Paragraph("一、问题背景与项目目标", h1),
            p_safe(
                "免费完整历史期权 tick 数据稀缺；学术研究中通常用「数据再现」方法：标的日线 + 实现波动率 + 合理 IV 曲面"
                "拼出虚拟历史期权链，再开展策略回测。本项目目标是闭环走通这套流程，并对结果做多维度风险归因。",
                body,
            ),
            p_safe("研究价值：", body),
            p_safe("· 卖方期权策略 (Short Strangle / Iron Condor) 的长期 P&L 分布特征", bullet),
            p_safe("· 风控参数 (止盈 / 止损 / 管理期) 对收益的边际贡献", bullet),
            p_safe("· 不同市场状态下损益的尾部特征与最大回撤来源", bullet),
            p_safe("· 组合层面的日度 Greeks 暴露与对冲需求", bullet),
        ]
    )

    # Section II
    story.extend(
        [
            Paragraph("二、历史期权数据再现方法", h1),
            Paragraph("2.1 数据流程", h2),
            p_safe(
                "(1) yfinance 拉取 SPY 三年日线；(2) 计算 20/60/120 日年化实现波动率 RV；"
                "(3) 每个历史交易日生成 5 个 DTE (7/14/30/45/60) 的虚拟期权链；(4) BS 模型 + IV 曲面定价。",
                body,
            ),
            Paragraph("2.2 IV 曲面公式", h2),
            p_safe(
                "IV(t, K, T) = max( 0.05,  RV_20(t) * IV_RV_RATIO  -  PUT_SKEW * ln(K / F_t) / sqrt(T) )",
                code,
            ),
            p_safe(f"· IV_RV_RATIO = {stats['iv_rv_ratio']:.2f}：观察到的 IV-RV 历史平均溢价", bullet),
            p_safe("· PUT_SKEW = 0.06：股票指数典型看跌偏度，使远端 OTM put 的 IV 高于 OTM call", bullet),
            p_safe("· 下层 0.05 截断：极端低波动期避免不合理的接近零的 IV", bullet),
        ]
    )
    add_image(PROJECT_OUT / "iv_rv_history.png", f"图1  回测期内 RV_20 与拟合 IV (IV = RV_20 x {stats['iv_rv_ratio']:.2f})", 16.0, 5.5, story)

    # Section III
    story.extend(
        [
            Paragraph("三、仿真交易策略与回测引擎", h1),
            Paragraph("3.1 策略规则", h2),
            p_safe(f"· 开仓：每周五，若空仓，按目标 delta = ±{stats['entry_delta']:.2f} 选 30 天到期的 call 与 put 构建宽跨卖方组合", bullet),
            p_safe("· 仓位规模：每笔 $50,000 现金担保，按行权价 × 100 计算合约数", bullet),
            p_safe(f"· 止盈：MTM ≥ {stats['take_profit']:.0%} 初始权利金时平仓", bullet),
            p_safe(f"· 止损：MTM ≤ -{stats['stop_loss']:.0%} 初始权利金时平仓", bullet),
            p_safe(f"· 管理期：DTE ≤ {stats['managed_dte']} 时强制平仓 (避开 gamma & pin risk)", bullet),
            p_safe("· 成本：每条腿 $0.65 佣金 × 开平合计 4 次；1% 中点价 bid-ask 滑点", bullet),
            Paragraph("3.2 引擎核心循环", h2),
            p_safe(
                "按日推进：1) 用当日 RV 重建 IV 曲面与当日虚拟期权链；2) 若持仓在线则重定价、检查风控三阈值；"
                "3) 若空仓且为周五则按规则开仓；4) 记录现金、MTM、组合 Greeks 与净值。",
                body,
            ),
        ]
    )

    # Section IV - results
    story.append(Paragraph("四、回测结果", h1))
    story.append(Paragraph("4.1 净值曲线与回撤", h2))
    add_image(PROJECT_OUT / "equity_curve.png", "图2  策略净值 vs SPY 买入持有", 16.0, 7.0, story)
    add_image(PROJECT_OUT / "drawdown.png", "图3  策略最大回撤序列", 16.0, 4.7, story)

    story.append(Paragraph("4.2 关键统计指标", h2))
    stat_rows = [
        ["指标", "数值", "指标", "数值"],
        ["初始资金", money(stats["start_equity"]), "Sharpe 比率", num(stats["sharpe_ratio"], 2)],
        ["结束净值", money(stats["end_equity"]), "Sortino 比率", num(stats["sortino_ratio"], 2)],
        ["总收益率", pct(stats["total_return"]), "最大回撤", pct(stats["max_drawdown"])],
        ["年化收益 (策略)", pct(stats["cagr"]), "胜率", pct(stats["win_rate"])],
        ["年化收益 (买入持有)", pct(stats["buyhold_cagr"]), "Profit Factor", num(stats["profit_factor"], 2)],
        ["交易次数", str(stats["n_trades"]), "期望 P&L / 笔", money(stats["expectancy"])],
        ["平均盈利", money(stats["avg_win"]), "平均亏损", money(stats["avg_loss"])],
        ["最大单笔盈利", money(stats["max_trade_pnl"]), "最大单笔亏损", money(stats["min_trade_pnl"])],
    ]
    story.append(std_table(stat_rows))
    story.append(Spacer(1, 5))

    story.append(Paragraph("4.3 交易级 P&L 分布", h2))
    add_image(PROJECT_OUT / "trade_pnl.png", "图4  左：交易 P&L 直方图   右：累计交易 P&L", 16.0, 6.0, story)

    story.append(Paragraph("4.4 组合 Greeks 暴露随时间变化", h2))
    add_image(PROJECT_OUT / "greek_exposure.png", "图5  日度 Delta / Gamma / Vega / Theta 暴露 (空仓时为 0)", 16.0, 12.0, story)

    # Section V - sample trades
    if not trades.empty:
        story.append(Paragraph("五、典型交易样本", h1))
        top = trades.nlargest(2, "pnl")
        bottom = trades.nsmallest(2, "pnl")
        median = trades.iloc[[len(trades) // 2 - 1, len(trades) // 2]] if len(trades) >= 4 else trades.head(2)
        sample = (
            pd.concat([top, bottom, median])
            .drop_duplicates(subset=["open_date", "expiry"])
            .sort_values("open_date")
        )
        tt = [["开仓", "平仓", "天数", "Put K", "Call K", "权利金", "P&L", "退出原因"]]
        for _, r in sample.iterrows():
            tt.append(
                [
                    str(r["open_date"]),
                    str(r["close_date"]),
                    str(int(r["days_held"])),
                    num(r["put_strike"], 1),
                    num(r["call_strike"], 1),
                    money(r["open_credit"]),
                    money(r["pnl"]),
                    str(r["exit_reason"]),
                ]
            )
        story.append(std_table(tt))
        story.append(Spacer(1, 4))

    # Section VI - discussion
    story.extend(
        [
            Paragraph("六、讨论与改进方向", h1),
            Paragraph("6.1 卖方策略的「拣硬币」困境", h2),
            p_safe(
                f"胜率 {stats['win_rate']:.1%} 但 Profit Factor 仅 {stats['profit_factor']:.2f}。"
                f"平均亏损 ({money(stats['avg_loss'])}) 是平均盈利 ({money(stats['avg_win'])}) 的 "
                f"{abs(stats['avg_loss'] / stats['avg_win']):.2f} 倍——这就是「在压路机前拣硬币」的经典效应。"
                "买入持有 SPY 在同期年化 "
                f"{stats['buyhold_cagr']:.2%}，是卖期权策略难以战胜的基准。",
                body,
            ),
            Paragraph("6.2 IV 曲面假设的局限", h2),
            p_safe(
                f"IV = RV_20 × {stats['iv_rv_ratio']:.2f} 是历史平均溢价，但在不同市场状态下溢价波动很大。"
                "VIX 飙升时实际 IV 远超此倍数，本仿真低估了对冲难度；低波期 IV-RV 溢价压缩到 0.95 左右，又高估了实际权利金。"
                "改进：将 IV_RV_RATIO 改为状态依赖的回归模型（参数依赖 VIX / SKEW / trend）。",
                body,
            ),
            Paragraph("6.3 进一步研究方向", h2),
            p_safe("· 对卖方策略加入 Delta 中性对冲（每日用标的期货回填 |Δ| > 阈值的暴露）", bullet),
            p_safe("· 用 IV 高分位过滤入场：只在 IV percentile > 50 时开仓", bullet),
            p_safe("· Iron Condor 替代 Short Strangle：以盈利换取尾部锁定", bullet),
            p_safe("· 参数最优化：entry_delta × take_profit × stop_loss 网格敏感性分析", bullet),
            p_safe("· 升级到真实数据：接入 Polygon Options / CBOE DataShop 比较合成与真实 IV 差异", bullet),
        ]
    )

    # Section VII - conclusion
    story.extend(
        [
            Paragraph("七、结论", h1),
            p_safe(
                "本项目完整闭环地实现了「数据再现 → 期权链合成 → 策略仿真 → 风控执行 → 风险/收益归因」整套流程。"
                f"SPY 样本中胜率 {stats['win_rate']:.1%} 的滚动短宽跨策略最终取得 {stats['total_return']:.2%} 总收益，"
                f"远低于买入持有 {stats['buyhold_total_return']:.2%}。"
                "结果不是策略失败，而是揭示了无对冲、无入场过滤的纯卖方策略在牛市中的系统性劣势。"
                "后续研究重点在 Delta 对冲、状态依赖入场、IV 曲面动态拟合上。",
                body,
            ),
            Paragraph("八、附件清单", h1),
            p_safe("· code/week10_project_backtest.py：完整回测引擎", bullet),
            p_safe("· code/generate_week10_project_report.py：本报告生成器", bullet),
            p_safe("· data/spy_history_backtest.csv, spy_history_with_rv.csv", bullet),
            p_safe("· outputs/equity_curve.csv, trades.csv, stats.csv, stats.json", bullet),
            p_safe("· outputs/equity_curve.png, drawdown.png, trade_pnl.png, greek_exposure.png, iv_rv_history.png", bullet),
        ]
    )

    doc.build(story)


def make_zip() -> Path:
    zip_path = ROOT / "作业提交包" / f"{SUBMISSION_NAME}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in PACKAGE_DIR.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(PACKAGE_PARENT))
    return zip_path


def main() -> None:
    stats = json.loads((PROJECT_OUT / "stats.json").read_text(encoding="utf-8"))
    ensure_pkg()
    copy_artifacts()
    write_readme(stats)
    build_docx()
    build_pdf()
    zip_path = make_zip()
    print(f"DOCX: {DOCX_PATH}")
    print(f"PDF:  {PDF_PATH}")
    print(f"ZIP:  {zip_path}")


if __name__ == "__main__":
    main()
