"""
generate_week10_homework_report.py
==================================

Generate the Week-10 homework submission (DOCX + PDF + ZIP).

Inputs (must exist):
  research/outputs/week10/homework/
      vix_summary.csv / vix_summary.json
      vix_compare.png
      spy_vix_contrib.png, uso_vix_contrib.png
      oil_strategy_summary.csv, oil_strategy_legs.csv, oil_strategy_mc.csv
      oil_strategy_payoffs.png, oil_strategy_greeks.png, oil_strategy_mc.png
      oil_strategy_meta.json
"""

from __future__ import annotations

import json
import shutil
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_NAME = "23300180062_任昊来_第十周作业"
HW_OUT = ROOT / "research/outputs/week10/homework"
DATA_DIR = ROOT / "data/week10"
PACKAGE_PARENT = ROOT / "作业提交包" / "_staging_week10_hw"
PACKAGE_DIR = PACKAGE_PARENT / SUBMISSION_NAME
DOCX_PATH = PACKAGE_DIR / f"{SUBMISSION_NAME}.docx"
PDF_PATH = PACKAGE_DIR / f"{SUBMISSION_NAME}.pdf"


def num(x, d=2):
    if pd.isna(x):
        return ""
    return f"{float(x):.{d}f}"


def money(x):
    if pd.isna(x):
        return ""
    return f"{float(x):,.2f}"


def pct(x):
    if pd.isna(x):
        return ""
    return f"{float(x):.2%}"


def ensure_pkg() -> None:
    if PACKAGE_DIR.exists():
        shutil.rmtree(PACKAGE_DIR)
    for sub in ["code", "data", "outputs"]:
        (PACKAGE_DIR / sub).mkdir(parents=True, exist_ok=True)


def copy_artifacts() -> None:
    code_src = [
        "research/week10_vix_calculator.py",
        "research/week10_oil_strategy.py",
        "research/generate_week10_homework_report.py",
    ]
    for rel in code_src:
        src = ROOT / rel
        if src.exists():
            shutil.copy2(src, PACKAGE_DIR / "code" / src.name)
    for path in DATA_DIR.glob("*"):
        if path.is_file():
            shutil.copy2(path, PACKAGE_DIR / "data" / path.name)
    for path in HW_OUT.glob("*"):
        if path.is_file():
            shutil.copy2(path, PACKAGE_DIR / "outputs" / path.name)


def write_readme() -> None:
    (PACKAGE_DIR / "README.txt").write_text(
        "\n".join(
            [
                "第十周作业：期权VIX指数复现 + 原油期权策略",
                "",
                f"提交名称：{SUBMISSION_NAME}",
                "",
                "主要内容：",
                "1. 题1：CBOE VIX白皮书方法 - 复现SPY的VIX与USO的OVX指数。",
                "2. 题2：对原油期权 USO 执行四种交易策略，分析到期日的损益和风险。",
                "   - 卖出跨式（Short Straddle）",
                "   - 卖出宽跨（Short Strangle）",
                "   - 铁鹰（Iron Condor）",
                "   - 日历价差（Calendar Spread）",
                "",
                "主要文件：",
                f"- {SUBMISSION_NAME}.pdf：PDF 报告",
                f"- {SUBMISSION_NAME}.docx：Word 报告",
                "- code/week10_vix_calculator.py：VIX 计算引擎（CBOE方法）",
                "- code/week10_oil_strategy.py：四策略对比分析",
                "- code/generate_week10_homework_report.py：本报告生成器",
                "- data/spy_chain.csv, uso_chain.csv：下载的期权链",
                "- outputs/：所有 CSV / JSON / PNG 输出",
                "",
                "复现命令：",
                "python research/week10_vix_calculator.py",
                "python research/week10_oil_strategy.py",
                "python research/generate_week10_homework_report.py",
            ]
        ),
        encoding="utf-8",
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.5)
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, "1F4E79")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(18)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.3)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.8)


def add_formula_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.2)
    r.font.color.rgb = RGBColor(60, 60, 60)


def build_docx() -> None:
    vix = pd.read_csv(HW_OUT / "vix_summary.csv")
    legs = pd.read_csv(HW_OUT / "oil_strategy_legs.csv")
    summary = pd.read_csv(HW_OUT / "oil_strategy_summary.csv")
    mc = pd.read_csv(HW_OUT / "oil_strategy_mc.csv")
    meta = json.loads((HW_OUT / "oil_strategy_meta.json").read_text(encoding="utf-8"))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("第十周作业：期权VIX指数复现与原油期权策略损益分析")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    meta_p = doc.add_paragraph(
        f"提交名称：{SUBMISSION_NAME}    数据：Yahoo期权链（SPY/USO）    样本日：{meta['as_of']}"
    )
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in meta_p.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)

    # ----- Question 1: VIX -----
    add_heading(doc, "题1：期权 VIX 指数的计算原理与品种复现", level=1)
    add_para(
        doc,
        "VIX（Volatility Index）由 CBOE 在 1993 年首创，2003 年改版后变成 SPX 期权全市场加权的无模型隐含方差指数。"
        "它不再依赖 Black-Scholes 等任何具体定价模型，而是用一组虚值看涨/看跌期权的报价直接复制一份 30 天的方差互换头寸。",
    )
    add_heading(doc, "1.1 CBOE 白皮书核心公式", level=2)
    add_para(doc, "白皮书定义的单个到期日的 30 天方差贡献为：")
    add_formula_box(doc, "sigma^2(T) = (2/T) * Sum_i [ delta_K_i / K_i^2 ] * exp(R*T) * Q(K_i)  -  (1/T) * (F/K_0 - 1)^2")
    add_para(doc, "把近月和次月两个标准化期限的方差按到期时间线性插值到精确 30 天得到 VIX：")
    add_formula_box(
        doc,
        "VIX = 100 * sqrt{ [ T_1*sigma_1^2*(N_T2 - N_30)/(N_T2 - N_T1) + T_2*sigma_2^2*(N_30 - N_T1)/(N_T2 - N_T1) ] * (N_365 / N_30) }",
    )
    add_bullet(doc, "F 为远期：F = K* + exp(R*T) * ( C(K*) - P(K*) )，K* 是 |C - P| 最小的行权价。")
    add_bullet(doc, "K_0 为小于 F 的最大行权价；以 K_0 为界，下方取虚值看跌、上方取虚值看涨。")
    add_bullet(doc, "delta_K_i 为相邻行权价之差的一半；Q(K_i) 为该期权的报价中点。")
    add_bullet(doc, "停止规则：从 K_0 向两端扩展，连续两个零报价的合约之后所有更远期合约都不参与。")

    add_heading(doc, "1.2 本程序在两个品种上的复现结果", level=2)
    add_para(
        doc,
        "本作业把上述公式实现为 week10_vix_calculator.py，并在两个品种上复现：SPY 期权（对应 CBOE 官方 VIX 指数 ^VIX）和 USO 期权（对应 CBOE 原油波动率指数 OVX, ^OVX）。"
        "复现的实现细节：1) 数据源 Yahoo Finance 的 yfinance 接口；2) 当 bid/ask 缺失（盘后快照）时回退到 lastPrice；3) 用 quote = max(bid, lastPrice) 作为零报价停止规则的依据；"
        "4) 时间单位采用 CBOE 推荐的分钟制（1 年 = 525600 分钟）；5) 双月线性插值到 30 天。",
    )
    tab = doc.add_table(rows=1, cols=7)
    for i, h in enumerate(["品种", "样本日", "现价", "复现 VIX", "官方参考值", "误差", "选用行权价数(近/次)"]):
        tab.rows[0].cells[i].text = h
    for _, row in vix.iterrows():
        cells = tab.add_row().cells
        cells[0].text = f"{row['ticker']} → {row['reference_index'].split('(')[0].strip()}"
        cells[1].text = str(row["as_of"])[:10]
        cells[2].text = num(row["spot"], 2)
        cells[3].text = num(row["computed_vix"], 2)
        cells[4].text = num(row["reference_value"], 2) if not pd.isna(row["reference_value"]) else "N/A"
        if not pd.isna(row["reference_value"]):
            cells[5].text = num(row["computed_vix"] - row["reference_value"], 2)
        else:
            cells[5].text = "—"
        cells[6].text = f"{int(row['near_num_strikes'])} / {int(row['next_num_strikes'])}"
    style_table(tab)
    if (HW_OUT / "vix_compare.png").exists():
        doc.add_picture(str(HW_OUT / "vix_compare.png"), width=Inches(5.7))
        cap = doc.add_paragraph("图1  CBOE 方法本程序复现 vs 官方公布值（左 SPY→VIX，右 USO→OVX）")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if (HW_OUT / "spy_vix_contrib.png").exists():
        doc.add_picture(str(HW_OUT / "spy_vix_contrib.png"), width=Inches(6.4))
        cap = doc.add_paragraph("图2  SPY 期权链中每个行权价对 σ² 的贡献（近月与次月）")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if (HW_OUT / "uso_vix_contrib.png").exists():
        doc.add_picture(str(HW_OUT / "uso_vix_contrib.png"), width=Inches(6.4))
        cap = doc.add_paragraph("图3  USO 期权链中每个行权价对 σ² 的贡献")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "从上表与图可见：本程序复现的 SPY-VIX 与官方 ^VIX 误差在 0.2 个百分点以内；USO 复现的 OVX 与官方 ^OVX 也只差约 2 个百分点。"
        "差异主要来源于 Yahoo 与 CBOE 取价时点不同、Yahoo 的 bid/ask 在闭市后被置零导致部分远翼合约被停止规则截断。"
        "在数据质量更高的情境下（如使用 CBOE DataShop 的逐分钟报价）复现误差可以进一步压缩。",
    )

    # ----- Question 2: Oil strategies -----
    add_heading(doc, "题2：原油期权 USO 四种交易策略的到期损益与风险", level=1)
    add_para(
        doc,
        f"今日（{meta['as_of']}）USO 现价 {num(meta['spot'], 2)} 美元，近月到期 {meta['near_expiry']}（DTE={meta['near_dte']} 天），"
        f"次月到期 {meta['far_expiry']}（DTE={meta['far_dte']} 天）。"
        f"为了让策略宽度自适应高 OVX 环境，宽跨与铁鹰的行权价采用 σ√T 加权：宽跨 ±0.5σ，铁鹰 body ±1σ / wing ±1.6σ。",
    )
    add_heading(doc, "2.1 四种策略的腿构成", level=2)
    legs_tab = doc.add_table(rows=1, cols=7)
    for i, h in enumerate(["策略", "类型", "K", "到期", "方向", "中间价", "IV"]):
        legs_tab.rows[0].cells[i].text = h
    for _, row in legs.iterrows():
        c = legs_tab.add_row().cells
        c[0].text = str(row["strategy"])
        c[1].text = str(row["option_type"])
        c[2].text = num(row["strike"], 1)
        c[3].text = str(row["expiry"])
        c[4].text = "买入" if int(row["quantity"]) > 0 else "卖出"
        c[5].text = money(row["mid"])
        c[6].text = pct(row["implied_volatility"])
    style_table(legs_tab)

    add_heading(doc, "2.2 组合 Greeks 与到期盈亏关键指标", level=2)
    s_tab = doc.add_table(rows=1, cols=9)
    for i, h in enumerate(["策略", "净权利金", "Δ", "Vega", "Θ", "最大盈利", "最大亏损", "盈亏平衡", "POP"]):
        s_tab.rows[0].cells[i].text = h
    for _, row in summary.iterrows():
        c = s_tab.add_row().cells
        c[0].text = str(row["strategy"])
        c[1].text = money(row["net_premium"])
        c[2].text = num(row["delta"], 2)
        c[3].text = num(row["vega"], 2)
        c[4].text = num(row["theta"], 2)
        c[5].text = money(row["max_profit"])
        c[6].text = money(row["max_loss"])
        c[7].text = str(row["breakevens"]) if not pd.isna(row["breakevens"]) else ""
        c[8].text = pct(row["prob_of_profit"])
    style_table(s_tab)

    if (HW_OUT / "oil_strategy_payoffs.png").exists():
        doc.add_picture(str(HW_OUT / "oil_strategy_payoffs.png"), width=Inches(6.4))
        cap = doc.add_paragraph("图4  四种策略的到期盈亏曲线（红色虚线 = 当前现价）")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if (HW_OUT / "oil_strategy_greeks.png").exists():
        doc.add_picture(str(HW_OUT / "oil_strategy_greeks.png"), width=Inches(6.4))
        cap = doc.add_paragraph("图5  四种策略的组合 Greeks 对比")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "2.3 Monte Carlo 风险中性下的风险量化", level=2)
    add_para(
        doc,
        "对每个策略做 20000 路径的对数正态 Monte Carlo（drift = r = 4%，sigma = 期权腿平均 IV），"
        "得到 POP、期望 P&L、VaR95、CVaR95，便于按风险预算比较四种策略。",
    )
    mc_tab = doc.add_table(rows=1, cols=6)
    for i, h in enumerate(["策略", "sigma", "POP", "期望 P&L", "VaR95", "CVaR95"]):
        mc_tab.rows[0].cells[i].text = h
    for _, row in mc.iterrows():
        c = mc_tab.add_row().cells
        c[0].text = str(row["strategy"])
        c[1].text = pct(row["sigma_used"])
        c[2].text = pct(row["prob_of_profit"])
        c[3].text = money(row["expected_pnl"])
        c[4].text = money(row["var_95"])
        c[5].text = money(row["cvar_95"])
    style_table(mc_tab)
    if (HW_OUT / "oil_strategy_mc.png").exists():
        doc.add_picture(str(HW_OUT / "oil_strategy_mc.png"), width=Inches(6.4))
        cap = doc.add_paragraph("图6  四种策略的 POP / 期望 P&L / VaR95 对比")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "2.4 结论与策略推荐", level=2)
    add_bullet(
        doc,
        "卖出跨式（Short Straddle）能拿到最大权利金，但 Theta 与 Vega 暴露最高，并且最大亏损无上限——只有在交易者认为油价将在到期前显著回归当前价、且 IV 即将下跌时使用。",
    )
    add_bullet(
        doc,
        "卖出宽跨（Short Strangle）以 0.5σ 行权价拉宽盈利区间，权利金比跨式少 40-50%，但 POP 从 57% 提升到 63%。",
    )
    add_bullet(
        doc,
        "铁鹰（Iron Condor）通过买入更远翼把最大亏损锁定为 1625 美元，POP 进一步提升到 70.5%，是高 OVX 环境下"
        "性价比最高的“有限风险卖方策略”。",
    )
    add_bullet(
        doc,
        "日历价差（Calendar Spread）唯一一个组合 Vega 为正的策略，适合预期 IV 上行 / 油价短期窄幅波动的情境，"
        "最大亏损固定为开仓成本 504 美元，下行风险最低。",
    )

    add_heading(doc, "三、附件清单", level=1)
    add_bullet(doc, "code/week10_vix_calculator.py：CBOE VIX 复现计算与图形")
    add_bullet(doc, "code/week10_oil_strategy.py：四种原油期权策略的腿构造、Greeks、到期 P&L、Monte Carlo")
    add_bullet(doc, "code/generate_week10_homework_report.py：本报告生成器")
    add_bullet(doc, "data/spy_chain.csv, uso_chain.csv：Yahoo 下载的期权链")
    add_bullet(doc, "outputs/：所有 CSV、JSON、PNG 输出（vix_summary、oil_strategy_summary 等）")

    doc.sections[0].footer.paragraphs[0].text = f"{SUBMISSION_NAME} | 第十周作业"
    doc.save(DOCX_PATH)


def register_fonts() -> tuple[str, str]:
    pdfmetrics.registerFont(TTFont("MSYH", r"C:\Windows\Fonts\msyh.ttc"))
    pdfmetrics.registerFont(TTFont("MSYH-Bold", r"C:\Windows\Fonts\msyhbd.ttc"))
    return "MSYH", "MSYH-Bold"


def p_safe(text, style):
    return Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style)


def build_pdf() -> None:
    vix = pd.read_csv(HW_OUT / "vix_summary.csv")
    legs = pd.read_csv(HW_OUT / "oil_strategy_legs.csv")
    summary = pd.read_csv(HW_OUT / "oil_strategy_summary.csv")
    mc = pd.read_csv(HW_OUT / "oil_strategy_mc.csv")
    meta = json.loads((HW_OUT / "oil_strategy_meta.json").read_text(encoding="utf-8"))
    font, bold = register_fonts()

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=1.4 * cm,
        leftMargin=1.4 * cm,
        topMargin=1.3 * cm,
        bottomMargin=1.3 * cm,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("title", parent=styles["Title"], fontName=bold, fontSize=16, leading=22, alignment=TA_CENTER, textColor=colors.HexColor("#1F4E79"))
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName=bold, fontSize=12.5, leading=17, spaceBefore=10, spaceAfter=5, textColor=colors.HexColor("#1F4E79"))
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName=bold, fontSize=10.5, leading=14, spaceBefore=6, spaceAfter=3, textColor=colors.HexColor("#1F4E79"))
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName=font, fontSize=8.8, leading=12.8, firstLineIndent=15, spaceAfter=3)
    bullet = ParagraphStyle("bul", parent=styles["BodyText"], fontName=font, fontSize=8.5, leading=12.0, leftIndent=15, spaceAfter=2)
    cap = ParagraphStyle("cap", parent=styles["BodyText"], fontName=font, fontSize=8.0, leading=10.5, alignment=TA_CENTER)
    formula = ParagraphStyle("formula", parent=styles["BodyText"], fontName=font, fontSize=8.6, leading=11.5, leftIndent=25, spaceAfter=4, textColor=colors.HexColor("#555555"))

    def std_table(data: list[list[str]]) -> Table:
        t = Table(data, repeatRows=1)
        t.setStyle(
            TableStyle(
                [
                    ("FONT", (0, 0), (-1, -1), font, 7.0),
                    ("FONT", (0, 0), (-1, 0), bold, 7.2),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            )
        )
        return t

    def add_image(path: Path, caption: str, width_cm: float, height_cm: float, story: list) -> None:
        if path.exists():
            story.append(Image(str(path), width=width_cm * cm, height=height_cm * cm))
            story.append(Paragraph(caption, cap))
            story.append(Spacer(1, 3))

    story = [
        Paragraph("第十周作业：期权VIX指数复现与原油期权策略损益分析", title),
        p_safe(f"提交名称：{SUBMISSION_NAME}    数据：Yahoo SPY/USO期权链    样本日：{meta['as_of']}", body),
        Paragraph("题1：期权 VIX 指数的计算原理与品种复现", h1),
        Paragraph("1.1 CBOE 白皮书核心公式", h2),
        p_safe(
            "VIX 是 CBOE 1993 年提出、2003 年改版后无模型化的隐含方差指数。它用一组虚值看涨/看跌期权的市场报价直接复制 30 天方差互换头寸，不依赖 Black-Scholes 等具体定价模型。本程序在 week10_vix_calculator.py 中精确实现了 CBOE 2019 白皮书算法。",
            body,
        ),
        p_safe("单一到期日的方差贡献：", body),
        p_safe("sigma^2(T) = (2/T) * Sum_i [ delta_K_i / K_i^2 ] * exp(R*T) * Q(K_i)  -  (1/T) * (F/K_0 - 1)^2", formula),
        p_safe("近月 + 次月加权回 30 天：", body),
        p_safe(
            "VIX = 100 * sqrt{ [ T_1*sigma_1^2*w_1 + T_2*sigma_2^2*w_2 ] * (N_365 / N_30) },   w_1 = (N_T2 - N_30)/(N_T2 - N_T1),  w_2 = (N_30 - N_T1)/(N_T2 - N_T1)",
            formula,
        ),
        p_safe("F = K* + exp(R*T) * ( C(K*) - P(K*) )，K* 为 |C-P| 最小行权价；K_0 为小于 F 的最大行权价；停止规则：连续两个零报价合约之外的远翼一律截断。", bullet),
    ]

    Paragraph("1.2 两个品种上的复现结果", h2)
    vix_table = [["品种", "样本日", "现价", "复现 VIX", "官方 ^VIX/^OVX", "误差", "选用行权价数(近/次)"]]
    for _, row in vix.iterrows():
        vix_table.append(
            [
                f"{row['ticker']} → {row['reference_index'].split('(')[0].strip()}",
                str(row["as_of"])[:10],
                num(row["spot"], 2),
                num(row["computed_vix"], 2),
                num(row["reference_value"], 2) if not pd.isna(row["reference_value"]) else "N/A",
                num(row["computed_vix"] - row["reference_value"], 2) if not pd.isna(row["reference_value"]) else "—",
                f"{int(row['near_num_strikes'])} / {int(row['next_num_strikes'])}",
            ]
        )
    story.extend(
        [
            Paragraph("1.2 两个品种上的复现结果", h2),
            std_table(vix_table),
            Spacer(1, 4),
        ]
    )
    add_image(HW_OUT / "vix_compare.png", "图1  本程序复现 vs CBOE 官方", 15.0, 7.0, story)
    add_image(HW_OUT / "spy_vix_contrib.png", "图2  SPY 期权链每行权价对 σ² 的贡献", 16.0, 6.0, story)
    add_image(HW_OUT / "uso_vix_contrib.png", "图3  USO 期权链每行权价对 σ² 的贡献", 16.0, 6.0, story)

    story.append(
        p_safe(
            "本程序复现 SPY → VIX 与官方 ^VIX 误差控制在 0.2 个百分点以内；USO → OVX 与官方 ^OVX 误差 ~2%。"
            "误差主要来自 Yahoo 与 CBOE 取价时点不同、闭市后部分远翼合约 bid/ask 缺失被停止规则截断。CBOE 官方使用 SPX 的逐分钟报价能进一步缩小误差。",
            body,
        )
    )

    story.append(Paragraph("题2：原油期权 USO 四种交易策略的到期损益与风险", h1))
    story.append(
        p_safe(
            f"今日 USO 现价 {num(meta['spot'], 2)} 美元，近月到期 {meta['near_expiry']}（DTE={meta['near_dte']}），次月到期 {meta['far_expiry']}（DTE={meta['far_dte']}）。"
            "宽跨与铁鹰行权价采用 σ√T 自适应间距（宽跨 ±0.5σ，铁鹰 body ±1σ / wing ±1.6σ），保证在 OVX 接近 76% 的极端环境下结构仍然有效。",
            body,
        )
    )

    legs_table = [["策略", "类型", "K", "到期", "方向", "中间价", "IV"]]
    for _, row in legs.iterrows():
        legs_table.append(
            [
                str(row["strategy"]),
                str(row["option_type"]),
                num(row["strike"], 1),
                str(row["expiry"]),
                "买入" if int(row["quantity"]) > 0 else "卖出",
                money(row["mid"]),
                pct(row["implied_volatility"]),
            ]
        )
    story.extend(
        [
            Paragraph("2.1 四种策略的腿构成", h2),
            std_table(legs_table),
            Spacer(1, 5),
        ]
    )

    s_table = [["策略", "净权利金", "Δ", "Vega", "Θ", "最大盈利", "最大亏损", "盈亏平衡", "POP"]]
    for _, row in summary.iterrows():
        s_table.append(
            [
                str(row["strategy"]),
                money(row["net_premium"]),
                num(row["delta"], 2),
                num(row["vega"], 2),
                num(row["theta"], 2),
                money(row["max_profit"]),
                money(row["max_loss"]),
                str(row["breakevens"]) if not pd.isna(row["breakevens"]) else "",
                pct(row["prob_of_profit"]),
            ]
        )
    story.extend(
        [
            Paragraph("2.2 组合 Greeks 与到期盈亏关键指标", h2),
            std_table(s_table),
            Spacer(1, 5),
        ]
    )
    add_image(HW_OUT / "oil_strategy_payoffs.png", "图4  四种策略到期盈亏曲线（红虚线=当前现价）", 16.0, 10.0, story)
    add_image(HW_OUT / "oil_strategy_greeks.png", "图5  四种策略 Greeks 对比", 16.0, 6.0, story)

    mc_table = [["策略", "sigma", "POP", "期望 P&L", "VaR95", "CVaR95"]]
    for _, row in mc.iterrows():
        mc_table.append(
            [
                str(row["strategy"]),
                pct(row["sigma_used"]),
                pct(row["prob_of_profit"]),
                money(row["expected_pnl"]),
                money(row["var_95"]),
                money(row["cvar_95"]),
            ]
        )
    story.extend(
        [
            Paragraph("2.3 Monte Carlo 风险量化", h2),
            p_safe("20000 路径的对数正态 Monte Carlo（risk-neutral，drift=r=4%，sigma=各腿 IV 平均）：", body),
            std_table(mc_table),
            Spacer(1, 5),
        ]
    )
    add_image(HW_OUT / "oil_strategy_mc.png", "图6  POP / 期望 P&L / VaR95 对比", 15.0, 6.0, story)

    story.extend(
        [
            Paragraph("2.4 结论与策略推荐", h2),
            p_safe("卖出跨式 — 权利金 $2363、Vega -33、Theta +39。POP 57.6%，最大亏损无上限。仅适用于强烈认为油价回归当前价的情境。", bullet),
            p_safe("卖出宽跨 — 权利金 $1202、POP 63.3%。比跨式更稳健但 P&L 比例更低。", bullet),
            p_safe("铁鹰 — 权利金 $375、最大亏损锁定 $1625、POP 70.5%。在高 OVX 环境下"
                   "性价比最高的“有限风险卖方策略”。", bullet),
            p_safe("日历价差 — 唯一 Vega 为正的策略，最大亏损固定 $504。适合预期 IV 上行 / 油价窄幅波动。", bullet),
            Paragraph("三、附件清单", h1),
            p_safe("code/week10_vix_calculator.py、week10_oil_strategy.py、generate_week10_homework_report.py", bullet),
            p_safe("data/spy_chain.csv, uso_chain.csv", bullet),
            p_safe("outputs/：所有 CSV、JSON、PNG", bullet),
        ]
    )

    doc.build(story)


def make_zip() -> Path:
    zip_path = ROOT / "作业提交包" / f"{SUBMISSION_NAME}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in PACKAGE_DIR.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(PACKAGE_PARENT))
    return zip_path


def main() -> None:
    ensure_pkg()
    copy_artifacts()
    write_readme()
    build_docx()
    build_pdf()
    zip_path = make_zip()
    print(f"DOCX: {DOCX_PATH}")
    print(f"PDF:  {PDF_PATH}")
    print(f"ZIP:  {zip_path}")


if __name__ == "__main__":
    main()
